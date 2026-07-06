import os
import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('KAVE Enterprise Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Intelligent Manufacturing Digital Twin - Dual-Pipeline Vision Engine & Chatbot')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Glossary
    add_heading(doc, '1. Glossary', 1)
    glossary = {
        'PaDiM': 'Patch Distribution Modeling. Used for real-time streaming anomaly detection via multivariate Gaussian distributions.',
        'PatchCore': 'A state-of-the-art anomaly detection algorithm used for high-precision, manual deep inspection using a K-Nearest Neighbors (KNN) memory bank.',
        'Redis Stream': 'A data structure in Redis acting as an append-only log. Used to stream high-frequency video frames from the edge camera to the Vision Engine.',
        'Mahalanobis distance': 'A measure of the distance between a point and a distribution. Used by PaDiM to calculate anomaly scores based on the learned Gaussian stats.',
        'WideResNet-50-2': 'The deep convolutional neural network backbone used to extract feature maps for both PaDiM and PatchCore.',
        'Digital Twin': 'A virtual representation of a physical manufacturing system (CNC machines, cameras, etc.) updated in real-time.',
        'Gemini 1.5 Flash': 'A lightweight, fast LLM by Google used as the core of the Chatbot agent to process natural language into SQL and Flux queries.'
    }
    for term, definition in glossary.items():
        p = doc.add_paragraph()
        p.add_run(term + ': ').bold = True
        p.add_run(definition)
    
    # 2. Requirements
    add_heading(doc, '2. Requirements', 1)
    add_heading(doc, 'Functional Requirements', 2)
    doc.add_paragraph('- Real-time Vision Inspection: The system must process camera streams via Redis at high FPS with sub-second latency.')
    doc.add_paragraph('- Manual Deep Inspection: The system must provide a REST endpoint to upload images and return an anomaly heatmap using a memory bank.')
    doc.add_paragraph('- AI Chatbot: The system must allow users to query production data using natural language.')
    doc.add_paragraph('- Dashboarding: A unified UI must visualize real-time camera feeds, historical defects, and simulation data.')
    
    add_heading(doc, 'Non-functional Requirements', 2)
    doc.add_paragraph('- Latency: Live stream processing must execute in < 50ms per frame.')
    doc.add_paragraph('- Scalability: The architecture must support decoupling through message brokers (Redis/Kafka) for multiple edge cameras.')
    doc.add_paragraph('- Maintainability: Docker containerization for all microservices.')
    
    # 3. Architecture & Design
    add_heading(doc, '3. Architecture & Design', 1)
    doc.add_paragraph('The system utilizes a microservice architecture built heavily around Docker, FastAPI, and Streamlit. Data is managed across three distinct database engines: PostgreSQL (relational business logic), InfluxDB (time-series sensor telemetry), and Redis (in-memory streaming).')
    
    add_heading(doc, 'Dual-Pipeline Logic', 2)
    doc.add_paragraph('The vision engine incorporates a Dual-Pipeline to balance speed and accuracy:')
    doc.add_paragraph('1. Real-Time (PaDiM): Features extracted by WideResNet are evaluated against a pre-computed Gaussian distribution using Mahalanobis distance. This guarantees < 50ms latency.')
    doc.add_paragraph('2. Manual/API (PatchCore): High-precision evaluation against a Coreset KNN memory bank. This takes ~500ms but yields superior localization.')
    
    add_heading(doc, 'System Diagrams (Mermaid.js)', 2)
    try:
        with open('UML_DIAGRAMS.md', 'r', encoding='utf-8') as f:
            uml_content = f.read()
        
        # Split by markdown headers
        sections = uml_content.split('## ')
        for section in sections[1:]:
            lines = section.split('\n')
            section_title = lines[0].strip()
            add_heading(doc, section_title, 3)
            
            code_block = []
            in_code = False
            desc_lines = []
            for line in lines[1:]:
                if line.startswith('```'):
                    in_code = not in_code
                    continue
                if in_code:
                    code_block.append(line)
                elif line.strip() != '':
                    desc_lines.append(line.strip())
                    
            if desc_lines:
                doc.add_paragraph(' '.join(desc_lines))
                
            if code_block:
                try:
                    p = doc.add_paragraph('\n'.join(code_block))
                    p.style = 'Intense Quote'
                except KeyError:
                    p = doc.add_paragraph('\n'.join(code_block))
    except Exception as e:
        doc.add_paragraph(f"Could not load UML_DIAGRAMS.md: {e}")
        
    # 4. API & Technical Reference
    add_heading(doc, '4. API & Technical Reference', 1)
    add_heading(doc, 'FastAPI Endpoints', 2)
    
    p = doc.add_paragraph()
    p.add_run('POST /predict (Vision Engine)').bold = True
    doc.add_paragraph('Description: Runs manual deep inspection (PatchCore) on an uploaded image.')
    doc.add_paragraph('Payload: multipart/form-data containing the image file.')
    doc.add_paragraph('Response JSON:')
    try:
        doc.add_paragraph('{\n  "filename": "nut_001.png",\n  "score": 1.4532,\n  "heatmap_b64": "<base64_encoded_png>",\n  "inference_ms": 482.1\n}', style='Intense Quote')
    except KeyError:
        doc.add_paragraph('{\n  "filename": "nut_001.png",\n  "score": 1.4532,\n  "heatmap_b64": "<base64_encoded_png>",\n  "inference_ms": 482.1\n}')
    
    p = doc.add_paragraph()
    p.add_run('GET /config (Vision Engine)').bold = True
    doc.add_paragraph('Description: Retrieves the current vision engine configuration.')
    
    p = doc.add_paragraph()
    p.add_run('POST /predict (Prediction Engine)').bold = True
    doc.add_paragraph('Description: Predicts CNC tool wear stage and confidence level using XGBoost.')
    
    # 5. Source Code Docs
    add_heading(doc, '5. Source Code Docs', 1)
    doc.add_paragraph('The codebase is strictly modularized into several domains:')
    doc.add_paragraph('- services/vision-engine/app.py: Core logic for dual-pipeline inference. Spawns a background thread `redis_consumer_worker` for stream monitoring, while serving standard FastAPI routes for synchronous REST calls.')
    doc.add_paragraph('- services/dashboard/src/components/chat_agent.py: Initializes the LangChain Agent with Gemini 1.5 Flash. Binds PostgreSQL and InfluxDB tool schemas allowing the LLM to query directly.')
    doc.add_paragraph('- data-processing/ai_training/train_padim.py: A script to traverse normal metal nut imagery, calculating mean and covariance matrices for fast runtime checking.')
    
    # 6. Testing & Deployment
    add_heading(doc, '6. Testing & Deployment', 1)
    add_heading(doc, 'Deployment (Docker Compose)', 2)
    doc.add_paragraph('1. Environment Variables: Clone `.env.example` to `.env` and inject standard passwords and Gemini API Keys.')
    doc.add_paragraph('2. Orchestration: Run `docker compose up -d --build` to instantiate all containers simultaneously. Networks are bridged internally so microservices communicate over DNS (e.g., `kave_db:5432`).')
    
    add_heading(doc, 'Testing Configuration', 2)
    doc.add_paragraph('Tests are primarily conducted via stream mockers. `vision-producer` simulates a camera by injecting base64-encoded JPEGs into Redis `camera_stream`. Performance validation dictates that end-to-end latency metrics are logged for each frame.')

    output_path = 'KAVE_Enterprise_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation generated successfully as {output_path}")

if __name__ == "__main__":
    main()
